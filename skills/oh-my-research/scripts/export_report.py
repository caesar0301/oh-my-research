#!/usr/bin/env python3
"""Render publication-safe Markdown chapters into DOCX or PDF.

This is a THIN, SPEC-DRIVEN renderer. All presentation decisions (title page,
fonts, colors, sizes, table of contents, headers/footers, chapter order,
front/back matter) come from an LLM-authored document spec — not from
hardcoded policy here. The script only:

  1. resolves the spec (JSON) merged over minimal defaults,
  2. assembles the chapter Markdown the agent wrote,
  3. runs a mechanical publication-safety scan,
  4. renders bytes (DOCX via python-docx, PDF via reportlab).

Author the spec at docs/<mode>/_document.json (or pass --spec). Generate a
starter with --emit-spec, then edit it per report/scenario/language.
"""

from __future__ import annotations

import argparse
import copy
import html
import itertools
import json
import os
import re
import sys
from collections.abc import Iterable
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

MODES = ("survey", "report", "manuscript", "brief")

INTERNAL_PATTERNS = {
    "internal material ID": re.compile(
        r"(?<![\w-])\[?(?:P|W|G|S|E)-\d{1,6}\]?(?![\w-])", re.IGNORECASE
    ),
    # Product attribution ("Powered by oh-my-research") is allowed chrome.
    # Block internal workflow jargon that must not appear in reader-facing prose.
    "workflow name": re.compile(
        r"\bOMR\b|Evidence-Deep|THINK mode|SYNTH mode", re.IGNORECASE
    ),
    "gate or QA name": re.compile(r"\bGate\s+[ABDL]\b|\bQA[12]\b", re.IGNORECASE),
    "private artifact path": re.compile(
        r"(?:docs/plans/|\.omr/|tree-state\.json)", re.IGNORECASE
    ),
    "internal evidence label": re.compile(
        r"evidence\s+(?:grade|boundary)|boundary-tagged|"
        r"(?:grade|boundary)\s*:\s*(?:proven|suggests|inferred)",
        re.IGNORECASE,
    ),
}

# Minimal, neutral defaults. Platform-specific faces are applied in
# apply_platform_font_defaults() so macOS/Linux/Windows pick local-friendly names.
DEFAULT_SPEC: dict[str, Any] = {
    "title": None,
    "subtitle": None,
    "author": "",
    "date": None,
    "page": {"size": "A4", "margin_mm": 22},
    "fonts": {
        "body": {"size": 10.5},
        "heading": {},
        "mono": {"size": 9},
        "pdf_latin": "Helvetica",
        "pdf_cjk": "STSong-Light",
    },
    "colors": {
        "heading": ["#1F4E79", "#2F5597", "#44546A"],
        "table_header_bg": "#D9EAF7",
    },
    "heading_sizes": [18, 14, 12],
    "line_spacing": 1.25,
    "cover": {
        "enabled": True,
        "elements": ["title", "subtitle", "author", "date"],
    },
    "toc": {"enabled": True, "title": None, "depth": 3},
    "header": {"enabled": True, "text": None},
    "footer": {"enabled": True, "page_numbers": True, "attribution": True},
    "attribution": {"enabled": True, "text": None},
    "chapters": {"order": None, "include": None, "exclude": []},
    "drop_first_h1_matching_title": True,
}

ATTRIBUTION_TEXT = "Powered by oh-my-research"
LANG_STRINGS = {
    "en": {
        "toc": "Table of Contents",
        "subtitle": "Deep Research Report",
    },
    "zh-CN": {
        "toc": "目录",
        "subtitle": "深度研究报告",
    },
    "zh-TW": {
        "toc": "目錄",
        "subtitle": "深度研究報告",
    },
    "zh-HK": {
        "toc": "目錄",
        "subtitle": "深度研究報告",
    },
    "ja": {
        "toc": "目次",
        "subtitle": "深度調査レポート",
    },
    "ko": {
        "toc": "목차",
        "subtitle": "심층 연구 보고서",
    },
    "de": {
        "toc": "Inhaltsverzeichnis",
        "subtitle": "Tiefgehender Forschungsbericht",
    },
    "fr": {
        "toc": "Table des matières",
        "subtitle": "Rapport de recherche approfondi",
    },
    "es": {
        "toc": "Índice",
        "subtitle": "Informe de investigación en profundidad",
    },
    "pt-BR": {
        "toc": "Sumário",
        "subtitle": "Relatório de pesquisa aprofundada",
    },
    "pt-PT": {
        "toc": "Índice",
        "subtitle": "Relatório de investigação aprofundada",
    },
    "ru": {
        "toc": "Содержание",
        "subtitle": "Глубокий исследовательский отчёт",
    },
    "ar": {
        "toc": "جدول المحتويات",
        "subtitle": "تقرير بحث معمق",
    },
    "it": {
        "toc": "Indice",
        "subtitle": "Report di ricerca approfondita",
    },
    "nl": {
        "toc": "Inhoudsopgave",
        "subtitle": "Diepgaand onderzoeksrapport",
    },
    "pl": {
        "toc": "Spis treści",
        "subtitle": "Pogłębiony raport badawczy",
    },
    "tr": {
        "toc": "İçindekiler",
        "subtitle": "Derin araştırma raporu",
    },
    "vi": {
        "toc": "Mục lục",
        "subtitle": "Báo cáo nghiên cứu chuyên sâu",
    },
    "th": {"toc": "สารบัญ", "subtitle": "รายงานวิจัยเชิงลึก"},
    "id": {
        "toc": "Daftar Isi",
        "subtitle": "Laporan penelitian mendalam",
    },
    "hi": {"toc": "विषय सूची", "subtitle": "गहन शोध रिपोर्ट"},
}


def language_family(language: str) -> str:
    return (language or "en").split("-", 1)[0].lower()


def is_cjk_language(language: str) -> bool:
    return language_family(language) in {"zh", "ja", "ko"}


# CJK ideographs, kana, hangul, and CJK punctuation. Latin/digits are excluded so
# they keep a proper Latin face instead of the half-width forms a CID font gives.
CJK_CHARS = (
    r"\u1100-\u11FF\u2E80-\u2EFF\u3000-\u303F\u3040-\u30FF\u3130-\u318F"
    r"\u31F0-\u31FF\u3400-\u4DBF\u4E00-\u9FFF\uA960-\uA97F\uAC00-\uD7FF"
    r"\uF900-\uFAFF\uFE30-\uFE4F\uFF00-\uFFEF"
)
CJK_RUN_RE = re.compile(f"[{CJK_CHARS}]+")

# Characters the PDF standard Latin fonts can actually draw (WinAnsiEncoding).
# Anything outside this set (arrows, ✓, ★, ①, CJK) must go to a broad-coverage
# CID font, otherwise reportlab silently substitutes the wrong glyph.
WINANSI_EXTRA = frozenset(
    "\u20ac\u201a\u0192\u201e\u2026\u2020\u2021\u02c6\u2030\u0160\u2039\u0152"
    "\u017d\u2018\u2019\u201c\u201d\u2022\u2013\u2014\u02dc\u2122\u0161\u203a"
    "\u0153\u017e\u0178"
)


def latin_safe(char: str) -> bool:
    point = ord(char)
    if char in "\n\r\t":
        return True
    if 0x20 <= point <= 0x7E or 0xA0 <= point <= 0xFF:
        return True
    return char in WINANSI_EXTRA


CJK_ONE_RE = re.compile(f"[{CJK_CHARS}]")

# Bullet glyphs per nesting level (PDF; DOCX uses its own list styles).
BULLET_MARKERS = ("\u2022", "\u25e6", "\u25aa")


# --------------------------------------------------------------------------- #
# Cross-platform fonts (macOS + Linux + Windows)
# --------------------------------------------------------------------------- #
def platform_family() -> str:
    name = sys.platform
    if name == "darwin":
        return "macos"
    if name.startswith("linux"):
        return "linux"
    if name in ("win32", "cygwin", "msys"):
        return "windows"
    return name


def font_search_dirs() -> list[Path]:
    """Standard font directories on macOS, Linux, and Windows (existing only)."""
    home = Path.home()
    windir = Path(os.environ.get("WINDIR", r"C:\Windows"))
    candidates = [
        # Linux / BSD
        Path("/usr/share/fonts"),
        Path("/usr/local/share/fonts"),
        Path("/usr/share/fonts/truetype"),
        Path("/usr/share/fonts/opentype"),
        Path("/usr/share/fonts/TTF"),
        home / ".local/share/fonts",
        home / ".fonts",
        # macOS
        Path("/System/Library/Fonts"),
        Path("/System/Library/Fonts/Supplemental"),
        Path("/Library/Fonts"),
        home / "Library/Fonts",
        # Windows
        windir / "Fonts",
        Path("C:/Windows/Fonts"),
    ]
    seen: set[Path] = set()
    out: list[Path] = []
    for path in candidates:
        try:
            resolved = path.resolve()
        except OSError:
            continue
        if resolved in seen or not path.is_dir():
            continue
        seen.add(resolved)
        out.append(path)
    return out


# Preferred TrueType/OpenType basenames for broad Unicode / symbol coverage.
_SYMBOL_FONT_NAMES = (
    "Arial Unicode.ttf",
    "ArialUnicode.ttf",
    "NotoSansSymbols2-Regular.ttf",
    "NotoSansSymbols-Regular.ttf",
    "NotoSans-Regular.ttf",
    "DejaVuSans.ttf",
    "FreeSans.ttf",
    "FreeSerif.ttf",
    "Symbola.ttf",
    "seguisym.ttf",
    "SegoeUISymbol.ttf",
    "Apple Symbols.ttf",
)

# Optional exact paths checked first (fast path when present).
_SYMBOL_FONT_HINTS = (
    "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    "/Library/Fonts/Arial Unicode.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
    "/usr/share/fonts/truetype/noto/NotoSansSymbols2-Regular.ttf",
    "/usr/share/fonts/opentype/noto/NotoSansSymbols2-Regular.otf",
    "/usr/local/share/fonts/DejaVuSans.ttf",
    "/usr/share/fonts/truetype/freefont/FreeSerif.ttf",
    "/usr/share/fonts/truetype/ancient-fonts/Symbola.ttf",
    "C:/Windows/Fonts/seguisym.ttf",
    "C:/Windows/Fonts/ARIALUNI.TTF",
    "/System/Library/Fonts/Apple Symbols.ttf",
)


# CJK TrueType/OpenType font basenames for PDF embedding (broad Unicode coverage).
# Order matters: prefer YaHei (雅黑) and SimSun (宋体) per zh-CN convention,
# then PingFang/Songti on macOS, Noto/Source Han on Linux.
_CJK_FONT_NAMES = (
    # Windows-first (雅黑 / 宋体) — best for zh-CN readers
    "msyh.ttc",
    "msyh.ttf",
    "MSYH.TTC",
    "MSYH.TTF",
    "simsun.ttc",
    "SIMSUN.TTC",
    "simhei.ttf",
    "SIMHEI.TTF",
    # macOS — SimSong (≈SimSun), Songti, PingFang, Heiti
    "SimSong.ttc",
    "Songti.ttc",
    "PingFang.ttc",
    "PingFang SC.ttf",
    "STHeiti Medium.ttc",
    "STHeiti Light.ttc",
    "Hiragino Sans GB.ttc",
    # Linux — Noto / Source Han
    "NotoSansCJKsc-Regular.otf",
    "NotoSansCJKsc-Regular.ttf",
    "NotoSansSC-Regular.ttf",
    "NotoSansSC-Regular.otf",
    "SourceHanSansSC-Regular.otf",
    "SourceHanSansSC-Regular.otf",
    "wqy-zenhei.ttc",
    "wqy-microhei.ttc",
)

# Optional exact CJK font paths checked first (fast path when present).
_CJK_FONT_HINTS = (
    # Windows
    "C:/Windows/Fonts/msyh.ttc",
    "C:/Windows/Fonts/MSYH.TTC",
    "C:/Windows/Fonts/msyh.ttf",
    "C:/Windows/Fonts/simsun.ttc",
    "C:/Windows/Fonts/SIMSUN.TTC",
    "C:/Windows/Fonts/simhei.ttf",
    "C:/Windows/Fonts/SIMHEI.TTF",
    # macOS — SimSong first (≈SimSun, has Bold variant)
    "/System/Library/AssetsV2/com_apple_MobileAsset_Font7/857d6c90171c328a4892c1492291d34e401d7f25.asset/AssetData/SimSong.ttc",
    "/System/Library/Fonts/Supplemental/Songti.ttc",
    "/System/Library/Fonts/STHeiti Medium.ttc",
    "/System/Library/Fonts/STHeiti Light.ttc",
    "/System/Library/Fonts/Hiragino Sans GB.ttc",
    "/System/Library/AssetsV2/com_apple_MobileAsset_Font7/3419f2a427639ad8c8e139149a287865a90fa17e.asset/AssetData/PingFang.ttc",
    # Linux
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/truetype/noto/NotoSansCJKsc-Regular.otf",
    "/usr/share/fonts/opentype/noto/NotoSansCJKsc-Regular.otf",
    "/usr/share/fonts/truetype/noto/NotoSansSC-Regular.ttf",
    "/usr/share/fonts/opentype/noto/NotoSansSC-Regular.otf",
    "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
    "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
)


def discover_symbol_font_paths() -> list[Path]:
    """Collect candidate symbol/Unicode fonts once (hints + directory scan)."""
    found: list[Path] = []
    seen: set[Path] = set()

    def add(path: Path) -> None:
        try:
            key = path.resolve()
        except OSError:
            key = path
        if key in seen or not path.is_file():
            return
        seen.add(key)
        found.append(path)

    for hint in _SYMBOL_FONT_HINTS:
        add(Path(hint))

    wanted = {name.lower() for name in _SYMBOL_FONT_NAMES}
    remaining = set(wanted) - {p.name.lower() for p in found}
    if remaining:
        for root in font_search_dirs():
            if not remaining:
                break
            try:
                for path in root.rglob("*"):
                    name = path.name.lower()
                    if path.is_file() and name in remaining:
                        add(path)
                        remaining.discard(name)
                        if not remaining:
                            break
            except OSError:
                continue
    return found


def register_symbol_font(chars: set[str]) -> str | None:
    """Embed the best available system font covering `chars` (macOS/Linux/Windows)."""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFError, TTFont

    if not chars:
        return None

    best: tuple[int, str] | None = None
    for path in discover_symbol_font_paths():
        font_name = re.sub(r"[^A-Za-z0-9]+", "", path.stem) or "SymbolFallback"
        if font_name[0].isdigit():
            font_name = "F" + font_name
        try:
            if font_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(font_name, str(path)))
            covered = pdfmetrics.getFont(font_name).face.charToGlyph
        except (TTFError, KeyError, ValueError):
            continue
        score = sum(1 for ch in chars if ord(ch) in covered)
        if best is None or score > best[0]:
            best = (score, font_name)
        if score == len(chars):
            break
    if best is None or best[0] == 0:
        return None
    name = best[1]
    pdfmetrics.registerFontFamily(
        name, normal=name, bold=name, italic=name, boldItalic=name
    )
    return name


def discover_cjk_font_paths() -> list[Path]:
    """Collect candidate CJK TrueType/OpenType fonts (hints + directory scan)."""
    found: list[Path] = []
    seen: set[Path] = set()

    def add(path: Path) -> None:
        try:
            key = path.resolve()
        except OSError:
            key = path
        if key in seen or not path.is_file():
            return
        seen.add(key)
        found.append(path)

    for hint in _CJK_FONT_HINTS:
        add(Path(hint))

    wanted = {name.lower() for name in _CJK_FONT_NAMES}
    remaining = set(wanted) - {p.name.lower() for p in found}
    if remaining:
        for root in font_search_dirs():
            if not remaining:
                break
            try:
                for path in root.rglob("*"):
                    name = path.name.lower()
                    if path.is_file() and name in remaining:
                        add(path)
                        remaining.discard(name)
                        if not remaining:
                            break
            except OSError:
                continue
    return found


def _ttc_subfont_index(path: Path, prefer_weight: int = 400) -> int:
    """Pick the TTC subfont closest to `prefer_weight` (OS/2 usWeightClass).

    Songti.ttc subfont 0 is Black (900); blindly using index 0 embeds a heavy
    face and makes every CJK glyph look bold. This scans subfont metadata to
    find the Regular (400) or nearest-lighter face, with a Bold (700) lookup
    done separately by the caller.
    """
    try:
        from fontTools.ttLib import TTCollection

        tc = TTCollection(str(path))
        best_idx, best_dist = 0, 10_000
        for i, f in enumerate(tc.fonts):
            weight = f["OS/2"].usWeightClass if "OS/2" in f else 400
            dist = abs(weight - prefer_weight)
            if dist < best_dist:
                best_idx, best_dist = i, dist
        return best_idx
    except (KeyError, ValueError, OSError, RuntimeError):
        return 0


def register_cjk_font(chars: set[str] | None = None) -> str | None:
    """Embed the best available CJK TTF font for PDF (covers chars if given).

    A real embedded TTF font is preferred over a non-embedded CID font
    (STSong-Light) because the latter renders as blank/tofu on viewers
    that lack the Adobe CJK pack. On macOS/Linux/Windows we pick a system
    CJK face and embed its glyphs so the PDF is self-contained.

    Returns the *normal* face name; a matching bold face (if the same TTC
    has a Bold/Heavy subfont) is registered under ``<name>-Bold`` so that
    ``<b>`` spans render with a real bold weight instead of fake-bold.
    """
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFError, TTFont

    best: tuple[int, str, Path, int] | None = None  # (score, name, path, subfont)
    for path in discover_cjk_font_paths():
        # Pick the Regular (400) subfont, not blindly index 0.
        sub_idx = _ttc_subfont_index(path, prefer_weight=400)
        font_name = re.sub(r"[^A-Za-z0-9]+", "", path.stem) or "CjkFallback"
        if font_name[0].isdigit():
            font_name = "F" + font_name
        try:
            if font_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(
                    TTFont(font_name, str(path), subfontIndex=sub_idx)
                )
            covered = pdfmetrics.getFont(font_name).face.charToGlyph
        except (TTFError, KeyError, ValueError):
            continue
        # Score by coverage of requested chars (if any), else by CJK coverage.
        if chars:
            score = sum(1 for ch in chars if ord(ch) in covered)
        else:
            score = sum(1 for cp in range(0x4E00, 0x9FFF, 64) if cp in covered)
        if best is None or score > best[0]:
            best = (score, font_name, path, sub_idx)
        if chars and score == len(chars):
            break
        if not chars and score > 0:
            break
    if best is None or best[0] == 0:
        return None
    name = best[1]
    path = best[2]

    # Try to register a genuine bold subfont (weight 600-800) from the same TTC.
    bold_name = f"{name}-Bold"
    if bold_name not in pdfmetrics.getRegisteredFontNames():
        bold_idx = _ttc_subfont_index(path, prefer_weight=700)
        if bold_idx != best[3]:  # different subfont exists
            try:
                pdfmetrics.registerFont(
                    TTFont(bold_name, str(path), subfontIndex=bold_idx)
                )
                pdfmetrics.registerFontFamily(
                    name,
                    normal=name,
                    bold=bold_name,
                    italic=name,
                    boldItalic=bold_name,
                )
                return name
            except (TTFError, KeyError, ValueError):
                pass
    # No bold subfont found: map bold to normal (fake-bold, but at least not
    # the wrong heavy weight).
    pdfmetrics.registerFontFamily(
        name, normal=name, bold=name, italic=name, boldItalic=name
    )
    return name


def platform_docx_fonts(language: str) -> tuple[str, str]:
    """Pick DOCX Latin + East Asian face names that usually exist on this OS.

    OOXML only stores names; the viewer substitutes if missing. Prefer faces
    commonly packaged on each platform so LibreOffice/Word look right locally.
    For zh-CN, prefer 雅黑 (YaHei) > 宋体 (SimSun/Songti) per convention.
    """
    family = language_family(language)
    plat = platform_family()

    if plat == "macos":
        latin = "Helvetica Neue"
        # macOS has no Microsoft YaHei; use SimSun (宋体) > Songti > PingFang
        east = {
            "zh": "SimSun",
            "ja": "Hiragino Sans",
            "ko": "Apple SD Gothic Neo",
        }.get(family, "SimSun")
    elif plat == "linux":
        latin = "DejaVu Sans"
        east = {
            "zh": "Noto Sans CJK SC",
            "ja": "Noto Sans CJK JP",
            "ko": "Noto Sans CJK KR",
        }.get(family, "Noto Sans CJK SC")
    elif plat == "windows":
        latin = "Calibri"
        # Windows: 雅黑 first, fallback 宋体
        east = {
            "zh": "Microsoft YaHei",
            "ja": "Yu Gothic",
            "ko": "Malgun Gothic",
        }.get(family, "Microsoft YaHei")
    else:
        latin, east = "Calibri", "Microsoft YaHei"
    return latin, east


def apply_platform_font_defaults(spec: dict[str, Any], language: str) -> dict[str, Any]:
    """Fill unset font fields with platform-appropriate defaults."""
    latin, east = platform_docx_fonts(language)
    fonts = spec.setdefault("fonts", {})
    for role in ("body", "heading"):
        slot = fonts.setdefault(role, {})
        slot.setdefault("latin", latin)
        slot.setdefault("eastasia", east)
    fonts.setdefault("mono", {}).setdefault(
        "latin", "Courier New" if platform_family() != "linux" else "DejaVu Sans Mono"
    )
    fonts.setdefault("pdf_latin", "Helvetica")
    fonts.setdefault(
        "pdf_cjk",
        {
            "zh": "STSong-Light",
            "ja": "HeiseiMin-W3",
            "ko": "HYSMyeongJo-Medium",
        }.get(language_family(language), "STSong-Light"),
    )
    return spec


def deep_merge(base: dict[str, Any], override: dict[str, Any]) -> dict[str, Any]:
    result = copy.deepcopy(base)
    for key, value in override.items():
        if isinstance(value, dict) and isinstance(result.get(key), dict):
            result[key] = deep_merge(result[key], value)
        else:
            result[key] = copy.deepcopy(value)
    return result


def lang_string(spec: dict[str, Any], language: str, key: str) -> str:
    table = (
        LANG_STRINGS.get(language)
        or LANG_STRINGS.get(language_family(language))
        or LANG_STRINGS["en"]
    )
    return table.get(key, LANG_STRINGS["en"][key])


def attribution_label(spec: dict[str, Any]) -> str | None:
    """Reader-facing product credit, or None when attribution is disabled."""
    attr = spec.get("attribution") or {}
    if not attr.get("enabled", True):
        return None
    text = attr.get("text")
    if text is None:
        return ATTRIBUTION_TEXT
    text = str(text).strip()
    return text or None


def footer_shows_attribution(spec: dict[str, Any]) -> bool:
    footer = spec.get("footer") or {}
    return bool(footer.get("enabled", True) and footer.get("attribution", True))


# --------------------------------------------------------------------------- #
# Chapter assembly
# --------------------------------------------------------------------------- #
def strip_frontmatter(text: str) -> str:
    if text.startswith("---\n"):
        end = text.find("\n---\n", 4)
        if end >= 0:
            return text[end + 5 :]
    return text


def default_sort_key(path: Path) -> tuple[int, str]:
    """Abstract/overview (00-*) leads unless the spec sets an explicit order."""
    name = path.stem
    return (0, name) if name.startswith("00-") else (1, name)


def chapter_files(root: Path, spec: dict[str, Any]) -> list[Path]:
    chapters_dir = root / "chapters"
    search_root = chapters_dir if chapters_dir.is_dir() else root
    available = {
        p.name: p
        for p in search_root.glob("*.md")
        if not p.name.startswith("_") and not p.name.startswith(".")
    }
    order = (spec.get("chapters") or {}).get("order")
    include = (spec.get("chapters") or {}).get("include")
    exclude = set((spec.get("chapters") or {}).get("exclude") or [])

    if order:
        selected: list[Path] = []
        for name in order:
            fname = name if name.endswith(".md") else f"{name}.md"
            if fname in available and fname not in {
                e if e.endswith(".md") else f"{e}.md" for e in exclude
            }:
                selected.append(available[fname])
        return selected

    names = list(available)
    if include:
        include_set = {n if n.endswith(".md") else f"{n}.md" for n in include}
        names = [n for n in names if n in include_set]
    exclude_set = {n if n.endswith(".md") else f"{n}.md" for n in exclude}
    names = [n for n in names if n not in exclude_set]
    return sorted((available[n] for n in names), key=default_sort_key)


def assemble(root: Path, spec: dict[str, Any]) -> tuple[str, list[tuple[str, str]]]:
    files = chapter_files(root, spec)
    if not files:
        raise SystemExit(f"No Markdown chapters found in {root}")
    chapters = [
        (p.name, strip_frontmatter(p.read_text(encoding="utf-8")).strip())
        for p in files
    ]
    combined = "\n\n".join(text for _, text in chapters).strip() + "\n"
    return combined, chapters


def validate_publication(text: str) -> list[str]:
    findings: list[str] = []
    for label, pattern in INTERNAL_PATTERNS.items():
        for match in pattern.finditer(text):
            line = text.count("\n", 0, match.start()) + 1
            findings.append(f"{label} at line {line}: {match.group(0)!r}")
    return findings


def infer_title(text: str) -> str | None:
    match = re.search(r"^#\s+(.+)$", text, re.MULTILINE)
    return clean_inline(match.group(1)) if match else None


def slugify(value: str) -> str:
    value = re.sub(r"[^\w\u3400-\u9fff]+", "-", value.strip().lower(), flags=re.UNICODE)
    return value.strip("-") or "research"


# --------------------------------------------------------------------------- #
# Inline math ($...$) — LaTeX rendered to Unicode with sub/superscripts
# --------------------------------------------------------------------------- #
MATH_INLINE_RE = re.compile(r"\$(?!\$)([^$\n]+?)\$")

# LaTeX control words → Unicode. Unknown commands fall back to their bare name.
_MATH_COMMANDS = {
    "left": "",
    "right": "",
    "equiv": "\u2261",
    "approx": "\u2248",
    "sim": "\u223c",
    "simeq": "\u2243",
    "cong": "\u2245",
    "propto": "\u221d",
    "neq": "\u2260",
    "ne": "\u2260",
    "leq": "\u2264",
    "le": "\u2264",
    "geq": "\u2265",
    "ge": "\u2265",
    "ll": "\u226a",
    "gg": "\u226b",
    "times": "\u00d7",
    "div": "\u00f7",
    "cdot": "\u00b7",
    "ast": "\u2217",
    "star": "\u22c6",
    "circ": "\u2218",
    "bullet": "\u2219",
    "pm": "\u00b1",
    "mp": "\u2213",
    "oplus": "\u2295",
    "ominus": "\u2296",
    "otimes": "\u2297",
    "odot": "\u2299",
    "to": "\u2192",
    "rightarrow": "\u2192",
    "longrightarrow": "\u27f6",
    "Rightarrow": "\u21d2",
    "implies": "\u21d2",
    "Leftarrow": "\u21d0",
    "leftarrow": "\u2190",
    "leftrightarrow": "\u2194",
    "Leftrightarrow": "\u21d4",
    "iff": "\u21d4",
    "mapsto": "\u21a6",
    "uparrow": "\u2191",
    "downarrow": "\u2193",
    "infty": "\u221e",
    "partial": "\u2202",
    "nabla": "\u2207",
    "forall": "\u2200",
    "exists": "\u2203",
    "nexists": "\u2204",
    "in": "\u2208",
    "notin": "\u2209",
    "ni": "\u220b",
    "subset": "\u2282",
    "subseteq": "\u2286",
    "supset": "\u2283",
    "supseteq": "\u2287",
    "cup": "\u222a",
    "cap": "\u2229",
    "setminus": "\u2216",
    "emptyset": "\u2205",
    "varnothing": "\u2205",
    "wedge": "\u2227",
    "land": "\u2227",
    "vee": "\u2228",
    "lor": "\u2228",
    "neg": "\u00ac",
    "lnot": "\u00ac",
    "sum": "\u2211",
    "prod": "\u220f",
    "int": "\u222b",
    "oint": "\u222e",
    "coprod": "\u2210",
    "angle": "\u2220",
    "perp": "\u22a5",
    "parallel": "\u2225",
    "mid": "\u2223",
    "top": "\u22a4",
    "bot": "\u22a5",
    "vdash": "\u22a2",
    "models": "\u22a8",
    "ldots": "\u2026",
    "dots": "\u2026",
    "cdots": "\u22ef",
    "vdots": "\u22ee",
    "ddots": "\u22f1",
    "langle": "\u27e8",
    "rangle": "\u27e9",
    "lceil": "\u2308",
    "rceil": "\u2309",
    "lfloor": "\u230a",
    "rfloor": "\u230b",
    "prime": "\u2032",
    "hbar": "\u210f",
    "ell": "\u2113",
    "Re": "\u211c",
    "Im": "\u2111",
    "aleph": "\u2135",
    "deg": "\u00b0",
    "quad": "\u2003",
    "qquad": "\u2003\u2003",
    "alpha": "\u03b1",
    "beta": "\u03b2",
    "gamma": "\u03b3",
    "delta": "\u03b4",
    "epsilon": "\u03b5",
    "varepsilon": "\u03b5",
    "zeta": "\u03b6",
    "eta": "\u03b7",
    "theta": "\u03b8",
    "vartheta": "\u03d1",
    "iota": "\u03b9",
    "kappa": "\u03ba",
    "lambda": "\u03bb",
    "mu": "\u03bc",
    "nu": "\u03bd",
    "xi": "\u03be",
    "omicron": "\u03bf",
    "pi": "\u03c0",
    "varpi": "\u03d6",
    "rho": "\u03c1",
    "varrho": "\u03f1",
    "sigma": "\u03c3",
    "varsigma": "\u03c2",
    "tau": "\u03c4",
    "upsilon": "\u03c5",
    "phi": "\u03d5",
    "varphi": "\u03c6",
    "chi": "\u03c7",
    "psi": "\u03c8",
    "omega": "\u03c9",
    "Gamma": "\u0393",
    "Delta": "\u0394",
    "Theta": "\u0398",
    "Lambda": "\u039b",
    "Xi": "\u039e",
    "Pi": "\u03a0",
    "Sigma": "\u03a3",
    "Upsilon": "\u03a5",
    "Phi": "\u03a6",
    "Psi": "\u03a8",
    "Omega": "\u03a9",
}


def _latex_to_unicode(expr: str) -> str:
    """Convert a LaTeX math fragment to plain Unicode (no scripts applied)."""
    s = expr.strip()
    # Font/structure wrappers: keep the inner text only.
    wrapper = re.compile(
        r"\\(?:text|mathrm|mathbf|mathbb|mathcal|mathsf|mathtt|mathit|"
        r"boldsymbol|operatorname)\s*\{([^{}]*)\}"
    )
    for _ in range(4):
        new = wrapper.sub(r"\1", s)
        if new == s:
            break
        s = new
    s = re.sub(r"\\frac\s*\{([^{}]*)\}\s*\{([^{}]*)\}", r"(\1)/(\2)", s)
    s = re.sub(r"\\sqrt\s*\{([^{}]*)\}", "\u221a(\\1)", s)
    s = s.replace("\\\\", " ")
    s = re.sub(r"\\[,;:!> ]", " ", s)
    s = re.sub(
        r"\\([A-Za-z]+)", lambda m: _MATH_COMMANDS.get(m.group(1), m.group(1)), s
    )
    s = s.replace("\\{", "{").replace("\\}", "}")
    return re.sub(r"[ \t]{2,}", " ", s)


def _split_italic(text: str) -> Iterable[tuple[str, bool]]:
    """Yield (chunk, italic) runs; ASCII letters (variables) render italic."""
    for match in re.finditer(r"[A-Za-z]+|[^A-Za-z]+", text):
        chunk = match.group(0)
        yield chunk, chunk[0].isalpha() and chunk.isascii()


def iter_math_runs(expr: str) -> Iterable[tuple[str, str, bool]]:
    """Yield (text, script, italic) where script is 'base' | 'sub' | 'sup'."""
    s = _latex_to_unicode(expr)
    i, n = 0, len(s)
    while i < n:
        ch = s[i]
        if ch in "_^":
            script = "sub" if ch == "_" else "sup"
            i += 1
            if i < n and s[i] == "{":
                depth, i, start = 1, i + 1, i + 1
                while i < n and depth:
                    if s[i] == "{":
                        depth += 1
                    elif s[i] == "}":
                        depth -= 1
                    if depth:
                        i += 1
                group = s[start:i]
                i += 1
            else:
                group = s[i] if i < n else ""
                i += 1
            for chunk, italic in _split_italic(group):
                yield chunk, script, italic
        elif ch in "{}":
            i += 1
        else:
            start = i
            while i < n and s[i] not in "_^{}":
                i += 1
            for chunk, italic in _split_italic(s[start:i]):
                yield chunk, "base", italic


def render_math_pdf(expr: str) -> str:
    """Reportlab markup for an inline math fragment (content pre-escaped)."""
    out: list[str] = []
    for text, script, italic in iter_math_runs(expr):
        if not text:
            continue
        piece = html.escape(text)
        if italic:
            piece = f"<i>{piece}</i>"
        if script == "sub":
            piece = f"<sub>{piece}</sub>"
        elif script == "sup":
            piece = f"<super>{piece}</super>"
        out.append(piece)
    return "".join(out)


def math_plain(expr: str) -> str:
    """Flatten math to plain Unicode (drops script positioning)."""
    return "".join(text for text, _, _ in iter_math_runs(expr))


def clean_inline(text: str) -> str:
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = MATH_INLINE_RE.sub(lambda m: math_plain(m.group(1)), text)
    text = re.sub(r"[*_`~]", "", text)
    return text.strip()


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def markdown_blocks(text: str) -> Iterable[dict[str, object]]:
    lines = text.splitlines()
    i = 0
    paragraph: list[str] = []
    # Ordinal per nesting level. Authors (and LLMs) commonly write every ordered
    # item as "1.", so the sequence is counted here instead of trusted verbatim.
    counters: dict[int, int] = {}

    def flush() -> Iterable[dict[str, object]]:
        nonlocal paragraph
        if paragraph:
            counters.clear()
            yield {"type": "paragraph", "text": " ".join(x.strip() for x in paragraph)}
            paragraph = []

    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            yield from flush()
            i += 1
            continue
        if line.startswith("<!--"):
            yield from flush()
            while i < len(lines) and "-->" not in lines[i]:
                i += 1
            i += 1
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            yield from flush()
            counters.clear()
            yield {
                "type": "heading",
                "level": len(heading.group(1)),
                "text": clean_inline(heading.group(2)),
            }
            i += 1
            continue
        if i + 1 < len(lines) and "|" in line and is_table_separator(lines[i + 1]):
            yield from flush()
            rows = [split_table_row(line)]
            i += 2
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                rows.append(split_table_row(lines[i]))
                i += 1
            counters.clear()
            yield {"type": "table", "rows": rows}
            continue
        bullet = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
        numbered = re.match(r"^(\s*)(\d+)[.)]\s+(.+)$", line)
        if bullet or numbered:
            yield from flush()
            indent = len((numbered or bullet).group(1).expandtabs(4))
            level = min(indent // 2, 2)
            for deeper in [key for key in counters if key > level]:
                del counters[deeper]
            block: dict[str, object] = {
                "type": "list",
                "ordered": bool(numbered),
                "level": level,
                "text": numbered.group(3) if numbered else bullet.group(2),
            }
            if numbered:
                # Honour an explicit start ("3." first) but renumber from there.
                block["first"] = level not in counters
                counters[level] = counters.get(level, int(numbered.group(2)) - 1) + 1
                block["index"] = counters[level]
            yield block
            i += 1
            continue
        if line.startswith(">"):
            yield from flush()
            counters.clear()
            yield {"type": "quote", "text": line.lstrip("> ").strip()}
            i += 1
            continue
        if line.startswith("```"):
            yield from flush()
            counters.clear()
            i += 1
            code: list[str] = []
            while i < len(lines) and not lines[i].startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            yield {"type": "code", "text": "\n".join(code)}
            continue
        paragraph.append(line)
        i += 1
    yield from flush()


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #
def _add_math_runs(
    paragraph: object,
    expr: str,
    latin: str | None = None,
    eastasia: str | None = None,
) -> None:
    for text, script, italic in iter_math_runs(expr):
        if not text:
            continue
        run = paragraph.add_run(text)
        if script == "sub":
            run.font.subscript = True
        elif script == "sup":
            run.font.superscript = True
        if italic:
            run.italic = True
        if latin and eastasia:
            _set_run_fonts(run, latin, eastasia)


def _docx_inline(
    paragraph: object,
    text: str,
    latin: str | None = None,
    eastasia: str | None = None,
    mono: str = "Courier New",
) -> None:
    token = re.compile(
        r"(\$(?!\$)[^$\n]+?\$|\*\*[^*]+\*\*|(?<!\*)\*[^*]+\*(?!\*)|`[^`]+`"
        r"|\[[^\]]+\]\([^)]+\))"
    )
    position = 0
    for match in token.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        value = match.group(0)
        if value.startswith("$"):
            _add_math_runs(paragraph, value[1:-1], latin, eastasia)
        elif value.startswith("**"):
            paragraph.add_run(value[2:-2]).bold = True
        elif value.startswith("*"):
            paragraph.add_run(value[1:-1]).italic = True
        elif value.startswith("`"):
            code_run = paragraph.add_run(value[1:-1])
            if latin and eastasia:
                _set_run_fonts(code_run, mono, eastasia)
            else:
                code_run.font.name = mono
        else:
            link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", value)
            run = paragraph.add_run(link.group(1) if link else value)
            if link:
                run.font.underline = True
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])
    if latin and eastasia:
        for run in paragraph.runs:
            if not _rfonts(run._element).get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii"
            ):
                _set_run_fonts(run, latin, eastasia)


def _rfonts(element: object) -> object:
    from docx.oxml import OxmlElement

    properties = element.get_or_add_rPr()
    fonts = properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        properties.insert(0, fonts)
    return fonts


def _set_fonts(element: object, latin: str, eastasia: str) -> None:
    """Bind Latin (ascii/hAnsi/cs) and East Asian faces separately.

    Without an explicit ascii/hAnsi face, Word renders Latin inside CJK text with
    the East Asian font, which mangles accents and letter spacing. An inherited
    w:hint="eastAsia" overrides ascii/hAnsi, so it is dropped as well.
    """
    from docx.oxml.ns import qn

    fonts = _rfonts(element)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:cs"), latin)
    fonts.set(qn("w:eastAsia"), eastasia)
    fonts.attrib.pop(qn("w:hint"), None)


def _set_run_fonts(run: object, latin: str, eastasia: str) -> None:
    _set_fonts(run._element, latin, eastasia)


def _set_paragraph_fonts(paragraph: object, latin: str, eastasia: str) -> None:
    for run in paragraph.runs:
        _set_run_fonts(run, latin, eastasia)


def _normalize_bullets(document: object, latin: str) -> None:
    """Replace Symbol-font private-use bullets with a real Unicode bullet.

    The default list template marks bullets as U+F0B7 in the Symbol font. When a
    renderer substitutes an East Asian font for that private-use codepoint it
    draws an arbitrary ideograph (e.g. 煉) instead of a bullet.
    """
    from docx.opc.exceptions import PackageNotFoundError
    from docx.oxml.ns import qn

    try:
        numbering = document.part.numbering_part.element
    except (AttributeError, KeyError, PackageNotFoundError, ValueError):
        return

    replacements = {"\uf0b7": "•", "\uf0a7": "▪", "\uf06c": "•", "\uf0d8": "‣"}
    for level in numbering.iter(qn("w:lvl")):
        text_el = level.find(qn("w:lvlText"))
        if text_el is None:
            continue
        value = text_el.get(qn("w:val")) or ""
        if not any(0xF000 <= ord(ch) <= 0xF0FF for ch in value):
            continue
        text_el.set(qn("w:val"), "".join(replacements.get(ch, "•") for ch in value))
        properties = level.find(qn("w:rPr"))
        if properties is not None:
            fonts = properties.find(qn("w:rFonts"))
            if fonts is not None:
                for attribute in ("w:ascii", "w:hAnsi", "w:cs"):
                    fonts.set(qn(attribute), latin)
                fonts.attrib.pop(qn("w:hint"), None)


def _fix_theme_fonts(document: object, latin: str, eastasia: str) -> None:
    """Set the theme's major/minor East Asian (and Latin) typefaces.

    python-docx ships a default theme whose <a:ea typeface=""/> is empty, so any
    style that references majorEastAsia/minorEastAsia (the built-in heading and
    Normal styles do) falls back to a viewer default — often the wrong CJK face
    or a blank box. Pinning the theme makes CJK render consistently everywhere.
    """
    from lxml import etree

    theme_part = None
    try:
        for part in document.part.package.iter_parts():
            if part.partname.endswith("/theme/theme1.xml"):
                theme_part = part
                break
    except (AttributeError, KeyError, ValueError):
        return
    if theme_part is None:
        return

    # The theme part is a plain Part (blob-based, not an XmlPart with _element).
    # Parse, patch, and write back to _blob so the change persists on save.
    try:
        root = etree.fromstring(theme_part.blob)
    except etree.XMLSyntaxError:
        return
    ns = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
    for font_tag in (f"{ns}majorFont", f"{ns}minorFont"):
        # majorFont/minorFont are nested under themeElements; use descendant search.
        for font_el in root.iter(font_tag):
            for child_tag in (f"{ns}latin", f"{ns}ea", f"{ns}cs"):
                child = font_el.find(child_tag)
                if child is None:
                    continue
                child.set("typeface", latin if child_tag != f"{ns}ea" else eastasia)
    theme_part._blob = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    )


def _style_num_id(document: object, style_id: str) -> str | None:
    """numId that a built-in list style points at, if any."""
    from docx.oxml.ns import qn

    for style in document.styles.element.findall(qn("w:style")):
        if style.get(qn("w:styleId")) == style_id:
            node = style.find(f"{qn('w:pPr')}/{qn('w:numPr')}/{qn('w:numId')}")
            return None if node is None else node.get(qn("w:val"))
    return None


def _restart_numbering(document: object, style_id: str, start: int = 1) -> int | None:
    """Clone a list style's numbering instance so a new list restarts at `start`.

    Every "List Number" paragraph shares one numId, so a second ordered list in
    the document keeps counting from the first instead of starting over.
    """
    from docx.opc.exceptions import PackageNotFoundError
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    source_id = _style_num_id(document, style_id)
    if source_id is None:
        return None
    try:
        numbering = document.part.numbering_part.element
    except (AttributeError, KeyError, PackageNotFoundError, ValueError):
        return None

    abstract = None
    used: set[int] = set()
    for num in numbering.findall(qn("w:num")):
        value = num.get(qn("w:numId")) or ""
        if value.isdigit():
            used.add(int(value))
        if value == source_id:
            abstract = num.find(qn("w:abstractNumId"))
    if abstract is None:
        return None

    new_id = (max(used) if used else 0) + 1
    element = OxmlElement("w:num")
    element.set(qn("w:numId"), str(new_id))
    reference = OxmlElement("w:abstractNumId")
    reference.set(qn("w:val"), abstract.get(qn("w:val")))
    element.append(reference)
    if start != 1:
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), str(start))
        override.append(start_override)
        element.append(override)
    numbering.append(element)
    return new_id


def _bind_numbering(paragraph: object, num_id: int) -> None:
    properties = paragraph._p.get_or_add_pPr()
    num_pr = properties.get_or_add_numPr()
    num_pr.get_or_add_ilvl().val = 0
    num_pr.get_or_add_numId().val = num_id


def _field(paragraph: object, instruction: str, placeholder: str = "") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    nodes = [begin, instr, sep]
    if placeholder:
        text_node = OxmlElement("w:t")
        text_node.text = placeholder
        sep.append(text_node)
    nodes.append(end)
    run._r.extend(nodes)


def export_docx(
    chapters: list[tuple[str, str]], output: Path, spec: dict[str, Any], language: str
) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
    except ImportError as exc:
        raise SystemExit(
            "DOCX export requires python-docx: pip install python-docx"
        ) from exc

    def rgb(hex_str: str) -> RGBColor:
        return RGBColor.from_string(hex_str.lstrip("#"))

    title = spec["title"]
    document = Document()
    section = document.sections[0]
    if str(spec["page"].get("size", "A4")).upper() == "A4":
        section.page_height, section.page_width = Cm(29.7), Cm(21)
    else:  # Letter
        section.page_height, section.page_width = Cm(27.94), Cm(21.59)
    margin = Cm(spec["page"].get("margin_mm", 22) / 10)
    section.top_margin = section.bottom_margin = margin
    section.left_margin = section.right_margin = margin

    body_font = spec["fonts"]["body"]["latin"]
    body_ea = spec["fonts"]["body"].get("eastasia", body_font)
    heading_font = spec["fonts"]["heading"].get("latin", body_font)
    heading_ea = spec["fonts"]["heading"].get("eastasia", body_ea)
    mono_font = spec["fonts"]["mono"]["latin"]

    normal = document.styles["Normal"]
    normal.font.name = body_font
    normal.font.size = Pt(spec["fonts"]["body"].get("size", 10.5))
    _set_fonts(normal._element, body_font, body_ea)
    normal.paragraph_format.line_spacing = spec.get("line_spacing", 1.25)
    normal.paragraph_format.space_after = Pt(7)
    _normalize_bullets(document, body_font)
    _fix_theme_fonts(document, body_font, body_ea)

    heading_colors = spec["colors"]["heading"]
    heading_sizes = spec["heading_sizes"]
    for level in range(1, 4):
        style = document.styles[f"Heading {level}"]
        style.font.name = heading_font
        _set_fonts(style._element, heading_font, heading_ea)
        style.font.color.rgb = rgb(
            heading_colors[min(level - 1, len(heading_colors) - 1)]
        )
        style.font.size = Pt(heading_sizes[min(level - 1, len(heading_sizes) - 1)])
        style.paragraph_format.space_before = Pt(14)
        style.paragraph_format.space_after = Pt(6)

    document.core_properties.title = title or ""
    document.core_properties.author = spec.get("author", "") or ""

    if spec["toc"].get("enabled", True):
        settings = document.settings.element
        if settings.find(qn("w:updateFields")) is None:
            update = OxmlElement("w:updateFields")
            update.set(qn("w:val"), "true")
            settings.append(update)

    # Cover (spec-driven element list)
    cover = spec.get("cover") or {}
    if cover.get("enabled", True):
        elements = cover.get("elements") or ["title"]
        date_str = spec.get("date") or datetime.now(timezone.utc).date().isoformat()
        first = True
        for element in elements:
            para = document.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if first:
                para.paragraph_format.space_before = Pt(150)
                first = False
            if element == "title" and title:
                run = para.add_run(title)
                run.bold = True
                run.font.size = Pt(28)
                run.font.color.rgb = rgb(heading_colors[0])
                _set_run_fonts(run, heading_font, heading_ea)
            elif element == "subtitle":
                text = spec.get("subtitle") or lang_string(spec, language, "subtitle")
                para.add_run(text).italic = True
            elif element == "author" and spec.get("author"):
                para.paragraph_format.space_before = Pt(60)
                para.add_run(spec["author"])
            elif element == "date":
                para.add_run(date_str)
            elif element == "attribution":
                credit = attribution_label(spec)
                if not credit:
                    continue
                para.paragraph_format.space_before = Pt(36)
                run = para.add_run(credit)
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = rgb("#666666")
            else:
                para.add_run(str(element))
            if element != "title":
                _set_paragraph_fonts(para, body_font, body_ea)
        document.add_page_break()

    # TOC
    if spec["toc"].get("enabled", True):
        toc_title = spec["toc"].get("title") or lang_string(spec, language, "toc")
        _set_paragraph_fonts(
            document.add_heading(toc_title, level=1), heading_font, heading_ea
        )
        depth = spec["toc"].get("depth", 3)
        _field(
            document.add_paragraph(),
            f' TOC \\o "1-{depth}" \\h \\z \\u ',
            "Update field to build the table of contents.",
        )
        document.add_page_break()

    drop_title = spec.get("drop_first_h1_matching_title", True)
    first_heading = True
    list_num_ids: dict[int, int | None] = {}
    for block in markdown_blocks("\n\n".join(text for _, text in chapters)):
        kind = block["type"]
        if kind == "heading":
            level = min(int(block["level"]), 3)
            text = str(block["text"])
            if first_heading and drop_title and level == 1 and title and text == title:
                first_heading = False
                continue
            first_heading = False
            _set_paragraph_fonts(
                document.add_heading(text, level=level), heading_font, heading_ea
            )
        elif kind == "paragraph":
            _docx_inline(
                document.add_paragraph(),
                str(block["text"]),
                body_font,
                body_ea,
                mono_font,
            )
        elif kind == "list":
            level = min(int(block.get("level", 0)), 2)
            suffix = "" if level == 0 else f" {level + 1}"
            style = f"{'List Number' if block['ordered'] else 'List Bullet'}{suffix}"
            paragraph = document.add_paragraph(style=style)
            if block["ordered"]:
                if block.get("first"):
                    list_num_ids[level] = _restart_numbering(
                        document,
                        style.replace(" ", ""),
                        int(block.get("index", 1)),
                    )
                if list_num_ids.get(level):
                    _bind_numbering(paragraph, list_num_ids[level])
            _docx_inline(
                paragraph,
                str(block["text"]),
                body_font,
                body_ea,
                mono_font,
            )
        elif kind == "quote":
            _docx_inline(
                document.add_paragraph(style="Quote"),
                str(block["text"]),
                body_font,
                body_ea,
                mono_font,
            )
        elif kind == "code":
            run = document.add_paragraph().add_run(str(block["text"]))
            _set_run_fonts(run, mono_font, body_ea)
            run.font.size = Pt(spec["fonts"]["mono"].get("size", 9))
        elif kind == "table":
            rows = block["rows"]
            if not rows:
                continue
            columns = max(len(r) for r in rows)
            table = document.add_table(rows=len(rows), cols=columns)
            table.style = "Light Shading Accent 1"
            for r_index, row in enumerate(rows):
                for c_index, value in enumerate(row):
                    cell = table.cell(r_index, c_index)
                    cell.text = clean_inline(value)
                    for run in cell.paragraphs[0].runs:
                        run.bold = r_index == 0
                        _set_run_fonts(run, body_font, body_ea)

    header_spec = spec.get("header") or {}
    footer_spec = spec.get("footer") or {}
    credit = attribution_label(spec) if footer_shows_attribution(spec) else None
    for section in document.sections:
        if header_spec.get("enabled", True):
            header_text = header_spec.get("text")
            header_text = title if header_text is None else header_text
            para = section.header.paragraphs[0]
            para.text = header_text or ""
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if para.runs:
                para.runs[0].font.size = Pt(8)
            _set_paragraph_fonts(para, body_font, body_ea)
        if footer_spec.get("enabled", True) and (
            footer_spec.get("page_numbers", True) or credit
        ):
            from docx.enum.text import WD_TAB_ALIGNMENT

            para = section.footer.paragraphs[0]
            para.text = ""
            if credit and footer_spec.get("page_numbers", True):
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = para.add_run(credit)
                run.font.size = Pt(8)
                run.font.color.rgb = rgb("#666666")
                _set_run_fonts(run, body_font, body_ea)
                para.add_run("\t")
                _field(para, " PAGE ")
                usable = section.page_width - section.left_margin - section.right_margin
                para.paragraph_format.tab_stops.add_tab_stop(
                    usable, WD_TAB_ALIGNMENT.RIGHT
                )
            elif credit:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(credit)
                run.font.size = Pt(8)
                run.font.color.rgb = rgb("#666666")
                _set_run_fonts(run, body_font, body_ea)
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _field(para, " PAGE ")

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


# --------------------------------------------------------------------------- #
# PDF
# --------------------------------------------------------------------------- #
class FontRouter:
    """Chooses a font per character class so mixed scripts render correctly.

    Latin stays on the Latin face: drawing a whole mixed paragraph with a CID font
    mangles accents ('ö') and letter spacing, while drawing symbols with a standard
    Latin font substitutes the wrong glyph or drops it silently.

    When a bold CJK variant is available, CJK runs inside <b> tags switch to it
    so bold Chinese text actually renders bold instead of staying regular.
    """

    def __init__(
        self,
        cjk_font: str | None,
        symbol_font: str | None,
        bold_cjk_font: str | None = None,
    ) -> None:
        self.cjk_font = cjk_font
        self.symbol_font = symbol_font
        self.bold_cjk_font = bold_cjk_font

    def __bool__(self) -> bool:
        return bool(self.cjk_font or self.symbol_font)

    def font_for(self, char: str, bold: bool = False) -> str | None:
        if latin_safe(char):
            return None
        if CJK_ONE_RE.match(char):
            if bold and self.bold_cjk_font:
                return self.bold_cjk_font
            return self.cjk_font or self.symbol_font
        return self.symbol_font or self.cjk_font

    def tag(self, markup: str) -> str:
        """Wrap non-Latin runs in <font> tags, tracking <b>/<i> context."""
        if not self:
            return markup
        out: list[str] = []
        bold = False
        for piece in re.split(r"(<[^>]+>)", markup):
            if not piece:
                continue
            if piece.startswith("<"):
                lower = piece.lower()
                if lower.startswith(("<b>", "<b ")):
                    bold = True
                elif lower.startswith("</b>"):
                    bold = False
                out.append(piece)
                continue
            def key(ch: str, b: bool = bold) -> str | None:
                return self.font_for(ch, bold=b)

            for font, group in itertools.groupby(piece, key=key):
                span = "".join(group)
                out.append(
                    span if font is None else f'<font name="{font}">{span}</font>'
                )
        return "".join(out)


def _pdf_inline(text: str, router: FontRouter | None = None) -> str:
    # Stash math first so escaping and markdown subs never touch its markup.
    math: list[str] = []

    def stash(match: re.Match[str]) -> str:
        math.append(render_math_pdf(match.group(1)))
        return f"\x00{len(math) - 1}\x00"

    value = MATH_INLINE_RE.sub(stash, text)
    value = html.escape(value)
    value = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", value)
    value = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"<i>\1</i>", value)
    value = re.sub(r"`([^`]+)`", r"<font name='Courier'>\1</font>", value)
    value = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"<a href='\2'>\1</a>", value)
    value = re.sub(r"\x00(\d+)\x00", lambda m: math[int(m.group(1))], value)
    return router.tag(value) if router else value


def export_pdf(
    chapters: list[tuple[str, str]], output: Path, spec: dict[str, Any], language: str
) -> None:
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
        from reportlab.lib.pagesizes import A4, LETTER
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.platypus import (
            BaseDocTemplate,
            Frame,
            PageBreak,
            PageTemplate,
            Paragraph,
            Spacer,
            Table,
            TableStyle,
        )
        from reportlab.platypus.tableofcontents import TableOfContents
    except ImportError as exc:
        raise SystemExit(
            "PDF export requires reportlab: pip install reportlab"
        ) from exc

    title = spec["title"]
    page_size = A4 if str(spec["page"].get("size", "A4")).upper() == "A4" else LETTER
    margin = spec["page"].get("margin_mm", 22) * mm

    standard_bold = {
        "Helvetica": "Helvetica-Bold",
        "Times-Roman": "Times-Bold",
        "Courier": "Courier-Bold",
    }
    font = spec["fonts"].get("pdf_latin", "Helvetica")
    bold_font = standard_bold.get(font, font)

    header_spec = spec.get("header") or {}
    footer_spec = spec.get("footer") or {}
    header_text = header_spec.get("text")
    header_text = title if header_text is None else header_text
    credit = attribution_label(spec) if footer_shows_attribution(spec) else None

    blocks = list(markdown_blocks("\n\n".join(text for _, text in chapters)))

    # Latin stays on a Latin face; CJK and symbol spans switch to a CID font.
    document_text = "\n".join(
        [
            str(title or ""),
            str(spec.get("subtitle") or ""),
            str(header_text or ""),
            str(credit or ""),
            *(text for _, text in chapters),
            # Injected list markers need coverage too: a CJK CID font has no
            # glyph for the nested bullets and would drop them silently.
            *(
                BULLET_MARKERS[min(int(block.get("level", 0)), len(BULLET_MARKERS) - 1)]
                for block in blocks
                if block["type"] == "list" and not block["ordered"]
            ),
            # Converted math symbols (≡ ∘ → …) are absent from the raw text, so
            # scan them explicitly or the CID font will silently drop them.
            *(
                math_plain(m)
                for _, text in chapters
                for m in MATH_INLINE_RE.findall(text)
            ),
        ]
    )
    non_latin = {ch for ch in document_text if not latin_safe(ch)}
    fallback_font: str | None = None
    if is_cjk_language(language) or any(CJK_ONE_RE.match(ch) for ch in non_latin):
        cjk_chars = {ch for ch in non_latin if CJK_ONE_RE.match(ch)}
        # Prefer an embedded TTF CJK font so the PDF is self-contained and
        # renders correctly on viewers without the Adobe CJK pack. Fall back
        # to a non-embedded CID font (STSong-Light) only if no TTF is found.
        embedded = register_cjk_font(cjk_chars or None)
        if embedded:
            fallback_font = embedded
        else:
            cjk_defaults = {
                "zh": "STSong-Light",
                "ja": "HeiseiMin-W3",
                "ko": "HYSMyeongJo-Medium",
            }
            candidate = spec["fonts"].get("pdf_cjk") or cjk_defaults.get(
                language_family(language), "STSong-Light"
            )
            for name in (candidate, "STSong-Light"):
                try:
                    pdfmetrics.registerFont(UnicodeCIDFont(name))
                    fallback_font = name
                    break
                except (KeyError, OSError):
                    continue
            if fallback_font:
                # CID fonts have no bold/italic face; map the family to itself so
                # <b>/<i> inside those spans do not raise during layout.
                pdfmetrics.registerFontFamily(
                    fallback_font,
                    normal=fallback_font,
                    bold=fallback_font,
                    italic=fallback_font,
                    boldItalic=fallback_font,
                )

    symbol_chars = {ch for ch in non_latin if not CJK_ONE_RE.match(ch)}
    # If a bold CJK variant was registered, pass it so <b>CJK</b> renders bold.
    bold_cjk = None
    if fallback_font:
        bold_candidate = f"{fallback_font}-Bold"
        if bold_candidate in pdfmetrics.getRegisteredFontNames():
            bold_cjk = bold_candidate
    router = FontRouter(
        fallback_font,
        register_symbol_font(symbol_chars) if symbol_chars else None,
        bold_cjk_font=bold_cjk,
    )

    def bullet_marker(level: int) -> str:
        """Nested bullet glyph, degraded to a WinAnsi bullet when uncovered."""
        glyph = BULLET_MARKERS[min(level, len(BULLET_MARKERS) - 1)]
        if latin_safe(glyph) or router.symbol_font:
            return glyph
        return BULLET_MARKERS[0]

    heading_colors = spec["colors"]["heading"]
    heading_sizes = spec["heading_sizes"]
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName=font,
        fontSize=spec["fonts"]["body"].get("size", 10.5),
        leading=spec["fonts"]["body"].get("size", 10.5) * 1.5,
        alignment=TA_JUSTIFY,
        spaceAfter=7,
        wordWrap="CJK" if is_cjk_language(language) else None,
    )

    def rich(text: str) -> str:
        """Escaped plain text with non-Latin spans bound to a covering font."""
        return router.tag(html.escape(text))

    headings = {
        level: ParagraphStyle(
            f"H{level}",
            parent=styles[f"Heading{level}"],
            fontName=bold_font,
            fontSize=heading_sizes[min(level - 1, len(heading_sizes) - 1)],
            leading=heading_sizes[min(level - 1, len(heading_sizes) - 1)] * 1.3,
            textColor=colors.HexColor(
                heading_colors[min(level - 1, len(heading_colors) - 1)]
            ),
            spaceBefore=13,
            spaceAfter=6,
        )
        for level in (1, 2, 3)
    }

    class ReportDoc(BaseDocTemplate):
        def afterFlowable(self, flowable: object) -> None:
            if isinstance(flowable, Paragraph) and flowable.style.name in (
                "H1",
                "H2",
                "H3",
            ):
                level = int(flowable.style.name[1]) - 1
                key = f"h-{self.seq.nextf('h')}"
                plain = flowable.getPlainText()
                self.canv.bookmarkPage(key)
                self.canv.addOutlineEntry(plain, key, level=level, closed=False)
                # TOC entries render as their own paragraphs, so they need the
                # same per-span font binding as the body.
                self.notify("TOCEntry", (level, rich(plain), self.page, key))

    def draw_mixed(canvas: object, x: float, y: float, text: str, size: float) -> None:
        """Draw chrome text span by span so each script keeps its own face.

        drawString takes a single face, so binding the whole line to a CID font
        would render Latin words such as the product name in CJK letterforms.
        """
        for name, group in itertools.groupby(text, key=router.font_for):
            span = "".join(group)
            face = name or font
            canvas.setFont(face, size)
            canvas.drawString(x, y, span)
            x += canvas.stringWidth(span, face, size)

    cover_enabled = bool((spec.get("cover") or {}).get("enabled", True))

    def decorate(canvas: object, doc: object) -> None:
        canvas.saveState()
        canvas.setFillColor(colors.HexColor("#666666"))
        # Skip header/footer chrome on the cover page (page 1) when a cover
        # is enabled — the cover is a clean title page, not a content page.
        is_cover = cover_enabled and doc.page == 1
        if not is_cover and header_spec.get("enabled", True) and header_text:
            draw_mixed(
                canvas, margin, page_size[1] - margin + 6 * mm, str(header_text)[:90], 8
            )
        y = margin - 8 * mm
        if footer_spec.get("enabled", True) and not is_cover:
            if credit:
                draw_mixed(canvas, margin, y, credit[:80], 7)
            if footer_spec.get("page_numbers", True):
                canvas.setFont(font, 8)
                canvas.drawRightString(page_size[0] - margin, y, str(doc.page))
        canvas.restoreState()

    output.parent.mkdir(parents=True, exist_ok=True)
    doc = ReportDoc(
        str(output),
        pagesize=page_size,
        title=title or "",
        author=spec.get("author", "") or "",
        leftMargin=margin,
        rightMargin=margin,
        topMargin=margin,
        bottomMargin=margin,
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
    doc.addPageTemplates([PageTemplate(id="pub", frames=frame, onPage=decorate)])

    story: list[object] = []
    cover = spec.get("cover") or {}
    if cover.get("enabled", True):
        cover_style = ParagraphStyle(
            "Cover",
            parent=styles["Title"],
            fontName=bold_font,
            fontSize=28,
            leading=35,
            alignment=TA_CENTER,
            textColor=colors.HexColor(heading_colors[0]),
        )
        center = ParagraphStyle(
            "Center", parent=body, alignment=TA_CENTER, fontName=font, fontSize=11
        )
        date_str = spec.get("date") or datetime.now(timezone.utc).date().isoformat()
        story.append(Spacer(1, 58 * mm))
        for element in cover.get("elements") or ["title"]:
            if element == "title" and title:
                story.append(Paragraph(rich(title), cover_style))
                story.append(Spacer(1, 12 * mm))
            elif element == "subtitle":
                story.append(
                    Paragraph(
                        rich(
                            spec.get("subtitle")
                            or lang_string(spec, language, "subtitle")
                        ),
                        center,
                    )
                )
                story.append(Spacer(1, 30 * mm))
            elif element == "author" and spec.get("author"):
                story.append(Paragraph(rich(spec["author"]), center))
            elif element == "date":
                story.append(Paragraph(date_str, center))
            elif element == "attribution":
                cover_credit = attribution_label(spec)
                if cover_credit:
                    story.append(Spacer(1, 18 * mm))
                    story.append(
                        Paragraph(
                            rich(cover_credit),
                            ParagraphStyle(
                                "Attribution",
                                parent=center,
                                fontSize=9,
                                textColor=colors.HexColor("#666666"),
                            ),
                        )
                    )
        story.append(PageBreak())

    if spec["toc"].get("enabled", True):
        toc_title = spec["toc"].get("title") or lang_string(spec, language, "toc")
        story.append(Paragraph(rich(toc_title), headings[1]))
        toc = TableOfContents()
        depth = spec["toc"].get("depth", 3)
        toc.levelStyles = [
            ParagraphStyle(
                f"TOC{level}",
                fontName=font,
                fontSize=max(11 - level, 8),
                leading=15,
                leftIndent=level * 12,
            )
            for level in range(depth)
        ]
        story.extend([toc, PageBreak()])

    drop_title = spec.get("drop_first_h1_matching_title", True)
    first_heading = True
    for block in blocks:
        kind = block["type"]
        if kind == "heading":
            level = min(int(block["level"]), 3)
            text = str(block["text"])
            if first_heading and drop_title and level == 1 and title and text == title:
                first_heading = False
                continue
            first_heading = False
            story.append(Paragraph(rich(text), headings[level]))
        elif kind == "paragraph":
            story.append(Paragraph(_pdf_inline(str(block["text"]), router), body))
        elif kind == "list":
            level = int(block.get("level", 0))
            marker = (
                f"{block.get('index', 1)}. "
                if block["ordered"]
                else bullet_marker(level) + " "
            )
            story.append(
                Paragraph(
                    router.tag(marker) + _pdf_inline(str(block["text"]), router),
                    ParagraphStyle(
                        f"List{level}",
                        parent=body,
                        leftIndent=14 + level * 16,
                        firstLineIndent=-10,
                    ),
                )
            )
        elif kind == "quote":
            story.append(
                Paragraph(
                    _pdf_inline(str(block["text"]), router),
                    ParagraphStyle(
                        "Quote",
                        parent=body,
                        leftIndent=14,
                        rightIndent=14,
                        textColor=colors.HexColor("#555555"),
                    ),
                )
            )
        elif kind == "code":
            # Apply the FontRouter so non-Latin chars (CJK, box-drawing, arrows)
            # inside code blocks switch to a covering font instead of rendering
            # as blank boxes when drawn purely in the Courier Latin face.
            code_html = router.tag(
                html.escape(str(block["text"])).replace("\n", "<br/>")
            )
            story.append(
                Paragraph(
                    code_html,
                    ParagraphStyle(
                        "Code",
                        parent=body,
                        fontName="Courier",
                        fontSize=8.5,
                        leading=11,
                        backColor=colors.HexColor("#F3F3F3"),
                        borderPadding=6,
                    ),
                )
            )
        elif kind == "table":
            rows = [
                [Paragraph(_pdf_inline(cell, router), body) for cell in row]
                for row in block["rows"]
            ]
            if not rows:
                continue
            widths = [doc.width / max(len(rows[0]), 1)] * len(rows[0])
            table = Table(rows, colWidths=widths, repeatRows=1, hAlign="LEFT")
            table.setStyle(
                TableStyle(
                    [
                        (
                            "BACKGROUND",
                            (0, 0),
                            (-1, 0),
                            colors.HexColor(
                                spec["colors"].get("table_header_bg", "#D9EAF7")
                            ),
                        ),
                        ("FONTNAME", (0, 0), (-1, 0), bold_font),
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#A6A6A6")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 5),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                        ("TOPPADDING", (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ]
                )
            )
            story.extend([table, Spacer(1, 6)])

    doc.multiBuild(story)


# --------------------------------------------------------------------------- #
# Spec resolution + CLI
# --------------------------------------------------------------------------- #
def resolve_spec(root: Path, args: argparse.Namespace, combined: str) -> dict[str, Any]:
    spec = copy.deepcopy(DEFAULT_SPEC)
    spec_path = args.spec or (root / "_document.json")
    if spec_path and Path(spec_path).exists():
        loaded = json.loads(Path(spec_path).read_text(encoding="utf-8"))
        spec = deep_merge(spec, loaded)
    # CLI overrides win over the file so quick runs still work
    if args.title:
        spec["title"] = args.title
    if args.author:
        spec["author"] = args.author
    if spec.get("title") is None:
        spec["title"] = infer_title(combined)
    apply_platform_font_defaults(spec, getattr(args, "language", None) or "en")
    return spec


def emit_spec(root: Path, language: str, mode: str) -> Path:
    spec = copy.deepcopy(DEFAULT_SPEC)
    apply_platform_font_defaults(spec, language)
    spec["title"] = None
    spec["_comment"] = (
        "LLM-authored document spec. Edit any field to control presentation. "
        "Delete fields to fall back to defaults. chapters.order lists chapter "
        "filenames (without .md) in publication order; omit to auto-order."
    )
    out = root / "_document.json"
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(
        json.dumps(spec, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
    )
    return out


def default_language(workspace: Path | None = None) -> str:
    """Timezone/locale-aware default; see references/LANGUAGE.md."""
    try:
        scripts_dir = Path(__file__).resolve().parent
        if str(scripts_dir) not in sys.path:
            sys.path.insert(0, str(scripts_dir))
        from prefer_language import resolve_language

        return resolve_language(workspace=workspace)["language"]
    except (ImportError, OSError, KeyError, ValueError):
        return "en"


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Spec-driven publication-safe report renderer"
    )
    parser.add_argument("--workspace", type=Path, default=Path.cwd())
    parser.add_argument("--mode", default="survey", choices=MODES)
    parser.add_argument("--format", default="docx", choices=["docx", "pdf"])
    parser.add_argument(
        "--language",
        default=None,
        help="Report language BCP-47 tag (default: timezone/workspace via prefer_language.py)",
    )
    parser.add_argument(
        "--spec", type=Path, default=None, help="Path to document spec JSON"
    )
    parser.add_argument(
        "--emit-spec",
        action="store_true",
        help="Write a starter _document.json and exit",
    )
    parser.add_argument("--title", default=None, help="Override spec title")
    parser.add_argument("--author", default="", help="Override spec author")
    parser.add_argument("--output", type=Path, default=None)
    args = parser.parse_args()

    workspace = args.workspace.resolve()
    language = args.language or default_language(workspace)
    root = workspace / "docs" / args.mode
    if args.emit_spec:
        print(emit_spec(root, language, args.mode))
        return
    if not root.is_dir():
        raise SystemExit(f"Missing docs/{args.mode}/")

    combined, chapters = assemble(root, {"chapters": {}})
    findings = validate_publication(combined)
    if findings:
        details = "\n".join(f"  - {f}" for f in findings[:30])
        raise SystemExit(
            "Publication-safety scan failed. Replace private references and workflow "
            f"terminology before export:\n{details}"
        )

    # Temporarily attach language onto args for resolve_spec consumers
    args.language = language
    spec = resolve_spec(root, args, combined)
    # Re-assemble honoring spec chapter order/include/exclude
    combined, chapters = assemble(root, spec)
    if spec.get("title") and validate_publication(spec["title"]):
        raise SystemExit("The report title contains private workflow terminology.")

    output = args.output
    if output is None:
        topic = slugify(spec.get("title") or args.mode)
        output = root / "deliverables" / f"{topic}-{args.mode}-{language}.{args.format}"
    elif not output.is_absolute():
        output = workspace / output

    if args.format == "docx":
        export_docx(chapters, output, spec, language)
    else:
        export_pdf(chapters, output, spec, language)
    print(output)


if __name__ == "__main__":
    main()
